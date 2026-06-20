# EN: Input agent for converting supported uploaded files into structured text.
# KO: 지원되는 업로드 파일을 구조화된 텍스트로 변환하는 Input Agent입니다.

from io import BytesIO
from typing import Any
import re
from datetime import date, datetime

from app.core.config import settings
from app.core.supported_files import (
    SUPPORTED_FILE_EXTENSIONS,
    SUPPORTED_FILE_TYPE_MESSAGE,
    resolve_supported_file_type,
)
from app.schemas.io_agent import (
    InputAgentRequest,
    InputAgentResponse,
    InputType,
    NormalizedRequestType,
)


class DocumentParserAgent:
    """Converts external document bytes into structured text for ingestion."""

    SUPPORTED_EXTENSIONS = SUPPORTED_FILE_EXTENSIONS
    PDF_PARSE_ERROR_MESSAGE = (
        "PDF 문서를 읽는 중 오류가 발생했습니다. 암호화된 PDF이거나 "
        "텍스트 추출이 불가능한 스캔본일 수 있습니다."
    )
    PDF_EMPTY_TEXT_MESSAGE = (
        "PDF에서 텍스트를 추출하지 못했습니다. 텍스트가 없는 스캔본 PDF일 수 있습니다."
    )
    EMPTY_TEXT_MESSAGE = (
        "문서에서 텍스트를 추출하지 못했습니다. 내용이 비어 있거나 텍스트 추출이 어려운 파일입니다."
    )
    AGENT_NAME = "DocumentParserAgent"

    async def parse(self, request: InputAgentRequest) -> InputAgentResponse:
        if request.input_type != InputType.FILE or not request.files:
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error="file input is required",
                validation_errors=["file input is required"],
            )

        file_payload = request.files[0]
        file_type = resolve_supported_file_type(
            file_name=file_payload.file_name,
            content_type=file_payload.content_type,
        )
        if file_type is None:
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error=SUPPORTED_FILE_TYPE_MESSAGE,
                validation_errors=[SUPPORTED_FILE_TYPE_MESSAGE],
            )
        extension = file_type.extension

        try:
            extraction_context: dict[str, Any] = {}
            if extension in {".xlsx", ".xls"}:
                text, extraction_context = self._extract_spreadsheet_text(
                    file_payload.file_bytes,
                    extension,
                    file_payload.file_name,
                )
            else:
                text = self._extract_text(file_payload.file_bytes, extension)
        except Exception as exc:
            error_message = self._parse_failure_message(extension, exc)
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error=error_message,
                validation_errors=[error_message],
            )

        text = self._remove_postgresql_unsafe_chars(text)
        if not text.strip():
            error_message = (
                self.PDF_EMPTY_TEXT_MESSAGE
                if extension == ".pdf"
                else self.EMPTY_TEXT_MESSAGE
            )
            return InputAgentResponse(
                success=False,
                agent_name=self.AGENT_NAME,
                normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
                error=error_message,
                validation_errors=[error_message],
            )

        metadata = {
            "file_name": file_payload.file_name,
            "extension": extension,
            "byte_size": len(file_payload.file_bytes),
            "content_type": file_payload.content_type,
        }
        if extraction_context.get("wbs_context"):
            metadata["wbs_context"] = extraction_context["wbs_context"]

        return InputAgentResponse(
            agent_name=self.AGENT_NAME,
            normalized_request_type=NormalizedRequestType.DOCUMENT_INGESTION,
            structured_context={
                "text": text,
                "metadata": metadata,
                **extraction_context,
            },
        )

    def _extract_text(self, file_bytes: bytes, extension: str) -> str:
        if extension == ".pdf":
            return self._extract_pdf_text(file_bytes)
        if extension == ".docx":
            return self._extract_docx_text(file_bytes)
        return self._decode_text(file_bytes)

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        errors: list[Exception] = []
        for module_name in ("pypdf", "PyPDF2"):
            try:
                module = __import__(module_name, fromlist=["PdfReader"])
            except ImportError:
                continue

            try:
                return self._extract_pdf_text_with_reader(
                    module.PdfReader,
                    file_bytes,
                )
            except Exception as exc:
                errors.append(exc)

        try:
            import pdfplumber
        except ImportError:
            pdfplumber = None
        if pdfplumber is not None:
            try:
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    self._ensure_pdf_page_limit(len(pdf.pages))
                    return "\n".join(
                        page.extract_text() or "" for page in pdf.pages
                    )
            except Exception as exc:
                errors.append(exc)

        try:
            import fitz
        except ImportError:
            fitz = None
        if fitz is not None:
            try:
                with fitz.open(stream=file_bytes, filetype="pdf") as document:
                    self._ensure_pdf_page_limit(getattr(document, "page_count", len(document)))
                    return "\n".join(page.get_text("text") or "" for page in document)
            except Exception as exc:
                errors.append(exc)

        if errors:
            raise errors[-1]
        raise RuntimeError("PDF parser library is not installed")

    def _extract_pdf_text_with_reader(self, reader_factory: Any, file_bytes: bytes) -> str:
        reader = reader_factory(BytesIO(file_bytes))
        if getattr(reader, "is_encrypted", False):
            decrypt = getattr(reader, "decrypt", None)
            if not callable(decrypt):
                raise ValueError("encrypted PDF")
            try:
                decrypt_result = decrypt("")
            except Exception as exc:
                raise ValueError("encrypted PDF") from exc
            if decrypt_result in (0, False):
                raise ValueError("encrypted PDF")

        pages = getattr(reader, "pages", [])
        self._ensure_pdf_page_limit(len(pages))
        return "\n".join(page.extract_text() or "" for page in pages)

    def _ensure_pdf_page_limit(self, page_count: int) -> None:
        max_pages = max(int(settings.UPLOAD_MAX_PDF_PAGES or 0), 1)
        if page_count > max_pages:
            raise ValueError(f"PDF page count exceeds limit: {page_count}/{max_pages}")

    def _parse_failure_message(self, extension: str, exc: Exception) -> str:
        if extension == ".pdf":
            if "parser library is not installed" in str(exc):
                return (
                    "PDF 문서를 읽기 위한 파서가 설치되어 있지 않습니다. "
                    "관리자에게 pypdf 설치를 요청해주세요."
                )
            return self.PDF_PARSE_ERROR_MESSAGE
        if extension == ".xls" and "xlrd is required" in str(exc):
            return (
                "XLS 문서를 읽기 위한 파서가 설치되어 있지 않습니다. "
                "관리자에게 xlrd 설치를 요청하거나 XLSX 형식으로 업로드해 주세요."
            )
        if extension in {".xlsx", ".xls"}:
            return (
                "엑셀 문서를 읽는 중 오류가 발생했습니다. 파일 형식과 시트 내용을 확인해주세요. "
                f"({type(exc).__name__})"
            )
        return f"문서를 읽는 중 오류가 발생했습니다. 파일 형식과 내용을 확인해주세요. ({type(exc).__name__})"

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for .docx uploads. Run `pip install python-docx`."
            ) from exc

        document = Document(BytesIO(file_bytes))
        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        numbering_levels = self._docx_numbering_levels(document)
        for table in document.tables:
            for row in table.rows:
                # sample_0605 preserved table column positions, including empty cells.
                # Removing empty cells shifts columns such as Biz요건ID/요구사항ID and
                # causes requirement extraction to map the wrong values.
                cells = [self._extract_docx_cell_text(cell, numbering_levels) for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))

        return "\n".join(lines)

    def _extract_docx_cell_text(self, cell, numbering_levels: dict[tuple[str, str], str]) -> str:
        paragraphs: list[str] = []
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            paragraphs.append(self._preserve_docx_list_marker(paragraph, text, numbering_levels))
        # Keep the table row as one physical line while preserving cell-internal
        # paragraph boundaries for downstream table extraction.
        return "\\n".join(paragraphs).strip()

    def _preserve_docx_list_marker(
        self,
        paragraph,
        text: str,
        numbering_levels: dict[tuple[str, str], str],
    ) -> str:
        if self._has_text_bullet_marker(text):
            return text
        list_level = self._docx_list_level(paragraph, numbering_levels)
        if list_level is None:
            return text
        marker = "-" if list_level > 0 else "o"
        return f"{marker} {text}"

    def _has_text_bullet_marker(self, text: str) -> bool:
        return bool(
            re.match(
                r"^\s*(?:[-*•ㅇ○·]|[oO]|\d+[.)]|[①②③④⑤⑥⑦⑧⑨⑩]|[ㄱ-ㅎ]\.)"
                r"(?:\s+|(?=[가-힣A-Z]))",
                text,
            )
        )

    def _is_docx_list_paragraph(self, paragraph) -> bool:
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
        if any(token in style_name for token in ["list", "bullet", "number"]):
            return True
        p_pr = getattr(getattr(paragraph, "_p", None), "pPr", None)
        if p_pr is not None and p_pr.numPr is not None:
            return True
        style = getattr(paragraph, "style", None)
        style_p_pr = getattr(getattr(style, "element", None), "pPr", None)
        return bool(style_p_pr is not None and style_p_pr.numPr is not None)

    def _docx_list_level(
        self,
        paragraph,
        numbering_levels: dict[tuple[str, str], str],
    ) -> int | None:
        p_pr = getattr(getattr(paragraph, "_p", None), "pPr", None)
        num_pr = getattr(p_pr, "numPr", None)
        num_id, ilvl_value = self._docx_num_id_and_level(num_pr)
        if ilvl_value is None:
            style = getattr(paragraph, "style", None)
            style_p_pr = getattr(getattr(style, "element", None), "pPr", None)
            style_num_pr = getattr(style_p_pr, "numPr", None)
            num_id, ilvl_value = self._docx_num_id_and_level(style_num_pr)

        if ilvl_value is not None:
            try:
                return int(ilvl_value)
            except (TypeError, ValueError):
                pass

        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
        match = re.search(r"(\d+)$", style_name)
        if match:
            return max(int(match.group(1)) - 1, 0)
        if self._is_docx_list_paragraph(paragraph):
            return 0
        return None

    def _docx_num_id_and_level(self, num_pr) -> tuple[str | None, str | None]:
        if num_pr is None:
            return None, None
        num_id = getattr(getattr(num_pr, "numId", None), "val", None)
        ilvl = getattr(getattr(num_pr, "ilvl", None), "val", None)
        return (str(num_id) if num_id is not None else None, str(ilvl) if ilvl is not None else None)

    def _docx_numbering_levels(self, document) -> dict[tuple[str, str], str]:
        levels: dict[tuple[str, str], str] = {}
        numbering_part = getattr(getattr(document, "part", None), "numbering_part", None)
        numbering = getattr(numbering_part, "element", None)
        if numbering is None:
            return levels

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        abstract_num_for_num: dict[str, str] = {}
        for num in numbering.findall("w:num", ns):
            num_id = self._xml_attr(num, "numId")
            abstract = num.find("w:abstractNumId", ns)
            abstract_id = self._xml_attr(abstract, "val")
            if num_id and abstract_id:
                abstract_num_for_num[num_id] = abstract_id

        abstract_levels: dict[tuple[str, str], str] = {}
        for abstract in numbering.findall("w:abstractNum", ns):
            abstract_id = self._xml_attr(abstract, "abstractNumId")
            if not abstract_id:
                continue
            for level in abstract.findall("w:lvl", ns):
                ilvl = self._xml_attr(level, "ilvl")
                fmt = level.find("w:numFmt", ns)
                num_fmt = self._xml_attr(fmt, "val") or ""
                if ilvl is not None:
                    abstract_levels[(abstract_id, ilvl)] = num_fmt

        for num_id, abstract_id in abstract_num_for_num.items():
            for (candidate_id, ilvl), num_fmt in abstract_levels.items():
                if candidate_id == abstract_id:
                    levels[(num_id, ilvl)] = num_fmt
        return levels

    def _xml_attr(self, element, name: str) -> str | None:
        if element is None:
            return None
        return element.get(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{name}")

    def _decode_text(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp949"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        return file_bytes.decode("utf-8", errors="ignore")

    def _extract_spreadsheet_text(
        self,
        file_bytes: bytes,
        extension: str,
        file_name: str | None,
    ) -> tuple[str, dict[str, Any]]:
        workbook_rows = (
            self._xlsx_rows(file_bytes)
            if extension == ".xlsx"
            else self._xls_rows(file_bytes)
        )
        lines: list[str] = []
        for sheet_name, rows in workbook_rows:
            for row_number, row in enumerate(rows, start=1):
                values = [value for value in row if value]
                if values:
                    lines.append(f"[{sheet_name} #{row_number}] " + " | ".join(values))

        wbs_context = self._extract_wbs_context_from_rows(
            workbook_rows,
            file_name=file_name,
        )
        context: dict[str, Any] = {}
        if wbs_context.get("rows"):
            context["wbs_context"] = wbs_context
        return "\n".join(lines), context

    def _xlsx_rows(self, file_bytes: bytes) -> list[tuple[str, list[list[str]]]]:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError(
                "openpyxl is required for .xlsx uploads. Run `pip install openpyxl`."
            ) from exc

        workbook = load_workbook(BytesIO(file_bytes), data_only=True, read_only=True)
        workbook_rows: list[tuple[str, list[list[str]]]] = []
        try:
            self._ensure_spreadsheet_sheet_limit(len(workbook.worksheets))
            for worksheet in workbook.worksheets:
                rows = []
                for row_number, row in enumerate(
                    worksheet.iter_rows(values_only=True),
                    start=1,
                ):
                    self._ensure_spreadsheet_row_limit(row_number)
                    rows.append([self._cell_to_text(value) for value in row])
                workbook_rows.append((worksheet.title, rows))
        finally:
            workbook.close()
        return workbook_rows

    def _xls_rows(self, file_bytes: bytes) -> list[tuple[str, list[list[str]]]]:
        try:
            import xlrd
        except ImportError as exc:
            raise RuntimeError(
                "xlrd is required for .xls uploads. Run `pip install xlrd`."
            ) from exc

        workbook = xlrd.open_workbook(file_contents=file_bytes)
        workbook_rows: list[tuple[str, list[list[str]]]] = []
        self._ensure_spreadsheet_sheet_limit(workbook.nsheets)
        for worksheet in workbook.sheets():
            self._ensure_spreadsheet_row_limit(worksheet.nrows)
            rows = [
                [self._cell_to_text(worksheet.cell_value(row_index, column_index)) for column_index in range(worksheet.ncols)]
                for row_index in range(worksheet.nrows)
            ]
            workbook_rows.append((worksheet.name, rows))
        return workbook_rows

    def _ensure_spreadsheet_sheet_limit(self, sheet_count: int) -> None:
        max_sheets = max(int(settings.UPLOAD_MAX_SPREADSHEET_SHEETS or 0), 1)
        if sheet_count > max_sheets:
            raise ValueError(
                f"spreadsheet sheet count exceeds limit: {sheet_count}/{max_sheets}"
            )

    def _ensure_spreadsheet_row_limit(self, row_count: int) -> None:
        max_rows = max(int(settings.UPLOAD_MAX_SPREADSHEET_ROWS or 0), 1)
        if row_count > max_rows:
            raise ValueError(
                f"spreadsheet row count exceeds limit: {row_count}/{max_rows}"
            )

    def _cell_to_text(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, datetime):
            return value.date().isoformat()
        if isinstance(value, date):
            return value.isoformat()
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value).strip()

    def _extract_wbs_context_from_rows(
        self,
        workbook_rows: list[tuple[str, list[list[str]]]],
        *,
        file_name: str | None,
    ) -> dict[str, Any]:
        rows: list[dict[str, Any]] = []
        for sheet_name, sheet_rows in workbook_rows:
            header_index, header_map = self._find_wbs_header(sheet_rows)
            if header_index is None:
                continue
            for row_offset, row in enumerate(sheet_rows[header_index + 1 :], start=header_index + 2):
                item = self._wbs_row_from_values(
                    row=row,
                    header_map=header_map,
                    row_number=row_offset,
                    file_name=file_name,
                    sheet_name=sheet_name,
                )
                if item:
                    rows.append(item)

        return {
            "source_document_name": file_name,
            "rows": rows,
        }

    def _find_wbs_header(
        self,
        rows: list[list[str]],
    ) -> tuple[int | None, dict[str, int]]:
        for index, row in enumerate(rows):
            normalized_cells = [self._normalize_header(cell) for cell in row]
            if not any(cell in {"wbs명", "작업명", "업무명", "taskname", "title"} for cell in normalized_cells):
                continue
            if not any(cell in {"시작예정일", "시작일", "plannedstartdate", "startdate"} for cell in normalized_cells):
                continue
            header_map = {
                normalized_cell: column_index
                for column_index, normalized_cell in enumerate(normalized_cells)
                if normalized_cell
            }
            return index, header_map
        return None, {}

    def _wbs_row_from_values(
        self,
        *,
        row: list[str],
        header_map: dict[str, int],
        row_number: int,
        file_name: str | None,
        sheet_name: str,
    ) -> dict[str, Any] | None:
        title = self._row_value(row, header_map, "wbs명", "작업명", "업무명", "taskname", "title", "name")
        if not title:
            return None

        level = self._row_value(row, header_map, "레벨", "level")
        wbs_id = self._row_value(row, header_map, "id", "wbsid", "wbs_id")
        planned_start = self._row_value(
            row,
            header_map,
            "시작예정일",
            "시작일",
            "plannedstartdate",
            "startdate",
        )
        planned_end = self._row_value(
            row,
            header_map,
            "종료예정일",
            "종료일",
            "plannedenddate",
            "enddate",
            "duedate",
        )
        assignee = self._row_value(row, header_map, "작업자", "담당자", "owner", "assignee")
        artifact = self._row_value(row, header_map, "산출물", "deliverable", "artifact")
        actual_start = self._row_value(row, header_map, "실제시작일", "actualstartdate")
        actual_end = self._row_value(row, header_map, "실제종료일", "actualenddate")
        status = self._row_value(row, header_map, "작업상태", "상태", "status")

        return {
            "row_number": row_number,
            "sheet_name": sheet_name,
            "level": level,
            "wbs_id": wbs_id,
            "title": title,
            "planned_start_date": planned_start,
            "planned_end_date": planned_end,
            "raw_assignee": assignee,
            "artifact": artifact,
            "actual_start_date": actual_start,
            "actual_end_date": actual_end,
            "raw_status": status,
            "source_document_name": file_name,
            "레벨": level,
            "ID": wbs_id,
            "WBS명": title,
            "시작예정일": planned_start,
            "종료예정일": planned_end,
            "작업자": assignee,
            "산출물": artifact,
            "실제시작일": actual_start,
            "실제종료일": actual_end,
            "작업상태": status,
        }

    def _row_value(
        self,
        row: list[str],
        header_map: dict[str, int],
        *aliases: str,
    ) -> str | None:
        for alias in aliases:
            column_index = header_map.get(self._normalize_header(alias))
            if column_index is None or column_index >= len(row):
                continue
            value = row[column_index].strip()
            if value:
                return value
        return None

    def _normalize_header(self, value: Any) -> str:
        return re.sub(r"[\s._/-]+", "", str(value or "").strip().lower())

    def _remove_postgresql_unsafe_chars(self, text: str) -> str:
        # PostgreSQL text/json columns cannot store NUL bytes. Binary Office files
        # decoded as plain text often contain them, so sanitize before persistence.
        return text.replace("\x00", "")


document_parser_agent = DocumentParserAgent()
